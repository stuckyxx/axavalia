import streamlit as st
import json
import os
from datetime import datetime, timedelta
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Removidos: from docx.oxml.ns import qn, from docx.oxml import OxmlElement
# Pois não serão mais usados para criar hyperlinks de baixo nível.

from docx2pdf import convert
import streamlit_authenticator as stauth
import yaml
from yaml.loader import SafeLoader

# --- CONFIGURAÇÕES E INICIALIZAÇÃO ---
if not os.path.exists("data/avaliacoes"):
    os.makedirs("data/avaliacoes")
if not os.path.exists("relatorios"):
    os.makedirs("relatorios")

# --- FUNÇÕES AUXILIARES ---
@st.cache_data
def carregar_criterios_do_arquivo(caminho_arquivo="criterios_por_topico.json"):
    """Carrega os critérios de avaliação e a lista de municípios do arquivo JSON."""
    try:
        with open(caminho_arquivo, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        st.error(f"ERRO: O arquivo de dados '{caminho_arquivo}' não foi encontrado.")
        return None
    except json.JSONDecodeError:
        st.error(f"ERRO: O arquivo '{caminho_arquivo}' contém um erro de formatação.")
        return None

def criar_pastas_necessarias():
    """Cria as pastas para salvar os dados e relatórios."""
    os.makedirs("data/avaliacoes", exist_ok=True)
    os.makedirs("relatorios", exist_ok=True)

def calcular_indice_e_selo(respostas, matriz_perguntas):
    """Calcula o índice de transparência e o selo Atricon com base nos pesos."""
    pesos = {"ESSENCIAL": 2.0, "OBRIGATÓRIA": 1.5, "RECOMENDADA": 1.0}
    total_pontos_possiveis, pontos_obtidos, total_essenciais, essenciais_atendidos = 0, 0, 0, 0
    for secao, perguntas in matriz_perguntas.items():
        if secao == "Municipios_MA": continue
        for item in perguntas:
            classificacao = item.get("classificacao", "RECOMENDADA").upper()
            peso = pesos.get(classificacao, 1.0)
            total_pontos_possiveis += peso
            status_geral_atende_item = True
            for sub in item["subcriterios"]:
                chave_resposta_sub = f"{secao}_{item['criterio']}_{sub}"
                if respostas.get(chave_resposta_sub) == "Não Atende":
                    status_geral_atende_item = False
                    break
            
            if status_geral_atende_item: pontos_obtidos += peso
            if classificacao == "ESSENCIAL":
                total_essenciais += 1
                if status_geral_atende_item: essenciais_atendidos += 1
                
    percentual_essenciais = (essenciais_atendidos / total_essenciais * 100) if total_essenciais > 0 else 100
    indice = (pontos_obtidos / total_pontos_possiveis * 100) if total_pontos_possiveis > 0 else 0
    
    selo = "Inexistente"
    if indice > 0:
        if percentual_essenciais == 100:
            if indice >= 95: selo = "💎 Diamante"
            elif indice >= 85: selo = "🥇 Ouro"
            elif indice >= 75: selo = "🥈 Prata"
            else: selo = "Elevado (não elegível para selo)"
        else:
            if indice >= 75: selo = "Elevado"
            elif indice >= 50: selo = "Intermediário"
            elif indice >= 30: selo = "Básico"
            else: selo = "Inicial"
    return {"indice": indice, "selo": selo, "percentual_essenciais": percentual_essenciais}


def calcular_pontuacao_secao(respostas, perguntas_secao, nome_secao):
    """Calcula a pontuação de uma seção específica."""
    pesos = {"ESSENCIAL": 2.0, "OBRIGATÓRIA": 1.5, "RECOMENDADA": 1.0}
    total_pontos_possiveis, pontos_obtidos = 0, 0
    for item in perguntas_secao:
        classificacao = item.get("classificacao", "RECOMENDADA").upper()
        peso = pesos.get(classificacao, 1.0)
        total_pontos_possiveis += peso
        
        item_atende = True
        for sub in item["subcriterios"]:
            chave_resposta = f"{nome_secao}_{item['criterio']}_{sub}"
            if respostas.get(chave_resposta) == "Não Atende":
                item_atende = False
                break
        
        if item_atende:
            pontos_obtidos += peso
    return (pontos_obtidos / total_pontos_possiveis * 100) if total_pontos_possiveis > 0 else 100

# Callback para quando a opção de Disponibilidade muda
def on_disponibilidade_change(secao, criterio, subcriterios):
    chave_disponibilidade = f"{secao}_{criterio}_Disponibilidade"
    novo_status_disponibilidade = st.session_state[chave_disponibilidade]
    st.session_state.respostas[chave_disponibilidade] = novo_status_disponibilidade

    if novo_status_disponibilidade == "Não Atende":
        for sub in subcriterios:
            if sub != "Disponibilidade":
                chave_subcriterio = f"{secao}_{criterio}_{sub}"
                st.session_state.respostas[chave_subcriterio] = "Não Atende"
                st.session_state.respostas[f"{chave_subcriterio}_obs"] = ""
    else:
        for sub in subcriterios:
            if sub != "Disponibilidade":
                chave_subcriterio = f"{secao}_{criterio}_{sub}"
                st.session_state.respostas[chave_subcriterio] = "Atende"
                st.session_state.respostas[f"{chave_subcriterio}_obs"] = ""
    st.rerun()

# --- FUNÇÃO DE GERAÇÃO DE RELATÓRIO ---
def gerar_relatorio_novo_modelo(respostas, municipio, segmento, matriz_perguntas, tipo_relatorio, nome_usuario, usuario_config):
    template_tipo = usuario_config.get('template', 'padrao')
    template_path = f"modelo_{template_tipo}.docx"
    
    if not os.path.exists(template_path):
        st.error(f"ERRO: Arquivo de modelo '{template_path}' não foi encontrado. Certifique-se de que ele está na mesma pasta do script.")
        return None, None 

    try:
        doc = docx.Document(template_path)
    except Exception as e:
        st.error(f"Erro ao carregar o modelo de relatório '{template_path}': {e}")
        return None, None
    
    # --- PÁGINA DE ROSTO ---
    for paragraph in doc.paragraphs:
        if "SEGMENTO" in paragraph.text:
            paragraph.text = paragraph.text.replace("SEGMENTO", segmento)
        if "NOME DO CLIENTE" in paragraph.text:
            paragraph.text = paragraph.text.replace("NOME DO CLIENTE", municipio)
        if "Data" in paragraph.text:
            paragraph.text = paragraph.text.replace("Data", datetime.now().strftime("%d/%m/%Y"))
    
    if template_tipo == 'padrao':
        doc.add_paragraph()
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("PROGRAMA NACIONAL DE TRANSPARÊNCIA PÚBLICA")
        run_title.font.size = Pt(22); run_title.bold = True

        p_subtitulo = doc.add_paragraph(); p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_subtitulo.add_run("Relatório de Transparência\n").bold = True
        p_subtitulo.add_run(f"{segmento} de {municipio}").bold = True
        
        doc.add_paragraph()
        resultados = calcular_indice_e_selo(respostas, matriz_perguntas)
        p_score = doc.add_paragraph(f"{resultados['indice']:.2f}%"); p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_score.runs[0].font.size = Pt(48); p_score.runs[0].bold = True
        p_selo = doc.add_paragraph(f"{resultados['selo']}"); p_selo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_selo.runs[0].font.size = Pt(24); p_selo.runs[0].bold = True
        doc.add_paragraph()
        
        texto_intro = f"Com base na Lei 12.527/2011 (Lei de Acesso à Informação), o nosso controle de qualidade fez uma avaliação geral da {segmento} de {municipio}, na qual, apresentou as seguintes informações:"
        doc.add_paragraph(texto_intro) # Não adiciona parágrafo extra aqui
        doc.add_paragraph() # Adiciona um parágrafo de espaço antes das infos de avaliação

        # Linhas de informação da avaliação com espaçamento
        p_exercicio = doc.add_paragraph(f"Exercício: {datetime.now().year}")
        doc.add_paragraph() # Espaço
        p_avaliador = doc.add_paragraph(f"Avaliação feita por: {nome_usuario}")
        doc.add_paragraph() # Espaço
        p_data_geracao = doc.add_paragraph(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph() # Espaço após a última linha

    doc.add_page_break()

    # --- PÁGINAS DE DETALHAMENTO ---
    p_detalhe = doc.add_paragraph()
    run_detalhe = p_detalhe.add_run("Detalhamento da Avaliação")
    run_detalhe.font.size = Pt(18); run_detalhe.bold = True
    doc.add_paragraph()

    matriz_para_relatorio = {}
    for secao, perguntas in matriz_perguntas.items():
        if secao == "Municipios_MA": continue

        itens_nao_conformes_na_secao = []
        for item in perguntas:
            algum_subcriterio_nao_atende = False
            for subcriterio in item["subcriterios"]:
                chave_resposta = f"{secao}_{item['criterio']}_{subcriterio}"
                if respostas.get(chave_resposta) == "Não Atende":
                    algum_subcriterio_nao_atende = True
                    break
            
            if algum_subcriterio_nao_atende:
                itens_nao_conformes_na_secao.append(item)
        
        if itens_nao_conformes_na_secao:
            matriz_para_relatorio[secao] = itens_nao_conformes_na_secao
    
    if not matriz_para_relatorio:
        doc.add_paragraph("Nenhuma não conformidade foi encontrada nesta avaliação. Todos os critérios foram atendidos.")
        doc.add_paragraph()
    else:
        for secao_nome, perguntas_secao_filtradas in matriz_para_relatorio.items():
            score_secao = calcular_pontuacao_secao(respostas, matriz_perguntas[secao_nome], secao_nome)
            
            p_secao_titulo = doc.add_paragraph()
            run_secao_titulo = p_secao_titulo.add_run(f"{secao_nome.upper()} - {score_secao:.2f}%")
            run_secao_titulo.font.size = Pt(14)
            run_secao_titulo.bold = True
            doc.add_paragraph()
            doc.add_paragraph()

            for item_idx, item in enumerate(perguntas_secao_filtradas):
                p_item_title = doc.add_paragraph()
                p_item_title.add_run(f"Item {item['topico']} - {item['criterio']} ({item.get('classificacao', '').upper()})")
                p_item_title.runs[0].bold = True
                doc.add_paragraph()

                observacoes_finais = []
                links_finais = []

                for subcriterio in item["subcriterios"]:
                    chave_resposta = f"{secao_nome}_{item['criterio']}_{subcriterio}"
                    resposta_sub = respostas.get(chave_resposta, "Atende")

                    if resposta_sub == "Não Atende":
                        p_criterio = doc.add_paragraph()
                        p_criterio.add_run(f"• {subcriterio}: ").italic = True
                        run_status = p_criterio.add_run("Não Atende")
                        run_status.bold = True
                        run_status.font.color.rgb = RGBColor(0xFF, 0, 0)
                        doc.add_paragraph()
                        
                        obs = respostas.get(f"{chave_resposta}_obs", "")
                        if obs: observacoes_finais.append((subcriterio, obs))

                chave_links_pergunta_geral = f"{secao_nome}_{item['criterio']}_links"
                links_gerais = respostas.get(chave_links_pergunta_geral, [])
                links_finais.extend(links_gerais)

                if links_finais or observacoes_finais:
                    p_obs_titulo = doc.add_paragraph()
                    p_obs_titulo.add_run("Evidências e Comentários:").bold = True
                    
                    for link_url in sorted(list(set(links_finais))):
                        if not (link_url.startswith("http://") or link_url.startswith("https://")):
                            link_url = "http://" + link_url
                        
                        p_obs_titulo.add_run(f"\n- Link: {link_url}")


                    for sub, obs_text in observacoes_finais:
                        p_obs_titulo.add_run(f"\n- Observação ({sub}): {obs_text}")
                    
                    doc.add_paragraph()
                    doc.add_paragraph()

    # --- SALVAMENTO E CONVERSÃO ---
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_base = f"Relatorio_Final_{segmento.replace(' ', '')}_{municipio.replace(' ', '')}_{timestamp}"
    path_docx = os.path.join("relatorios", f"{nome_base}.docx")
    path_pdf = os.path.join("relatorios", f"{nome_base}.pdf")
    
    doc.save(path_docx)
    try:
        convert(path_docx, path_pdf)
        return path_docx, path_pdf
    except Exception as e:
        st.error(f"Falha ao converter DOCX para PDF: {e}. O arquivo DOCX foi salvo e está disponível para download.")
        return path_docx, None


# --- INTERFACE GRÁFICA ---
st.set_page_config(layout="wide", page_title="Avaliador de Transparência")
st.title("📄 Sistema de Avaliação de Transparência Municipal")
matriz_completa = carregar_criterios_do_arquivo()

if matriz_completa:
    try:
        with open('config.yaml', 'r', encoding='utf-8') as file: config = yaml.load(file, Loader=SafeLoader)
        authenticator = stauth.Authenticate(
            config['credentials'],
            config['cookie']['name'],
            config['cookie']['key'],
            config['cookie']['expiry_days']
        )
        authenticator.login('main')

    except FileNotFoundError:
        st.error("ERRO: O arquivo 'config.yaml' não foi encontrado. Por favor, crie-o na raiz do projeto.")
        st.stop()
    except Exception as e:
        st.error(f"Erro ao carregar a configuração de autenticação: {e}. Verifique o formato do 'config.yaml'.")
        st.stop()

    if st.session_state["authentication_status"]:
        authenticator.logout('Logout', 'sidebar', key='logout_button')
        st.sidebar.title(f"Bem-vindo(a),\n{st.session_state['name']}!")
        
        st.sidebar.header("Configuração da Avaliação")
        
        MUNICIPIOS_MARANHAO = ["- Selecione um município -"] + sorted(matriz_completa.get("Municipios_MA", []))
        municipio = st.sidebar.selectbox("Nome do Município", options=MUNICIPIOS_MARANHAO, key="select_municipio")
        
        opcoes_segmento = [key for key in matriz_completa.keys() if key != "Municipios_MA"]
        segmento = st.sidebar.selectbox("Órgão/Poder", opcoes_segmento, key="select_segmento")
        
        if municipio != "- Selecione um município -" and segmento:
            nome_arquivo_avaliacao = f"avaliacao_{segmento.replace(' ', '')}_{municipio.replace(' ', '')}_{st.session_state['username']}.json"
            caminho_arquivo = os.path.join("data/avaliacoes", nome_arquivo_avaliacao)
            
            if st.sidebar.button("✅ Iniciar / Continuar Avaliação"):
                if os.path.exists(caminho_arquivo):
                    with open(caminho_arquivo, 'r', encoding='utf-8') as f: 
                        st.session_state.respostas = json.load(f)
                    st.sidebar.success("Avaliação anterior carregada!")
                else:
                    st.session_state.respostas = {}
                    st.sidebar.info("Iniciando uma nova avaliação.")
                
                st.session_state.path_pdf = None
                st.session_state.fallback_docx_path = None

                st.session_state.avaliacao_iniciada = True
                st.session_state.caminho_arquivo = caminho_arquivo
                st.session_state.municipio = municipio
                st.session_state.segmento = segmento
                st.session_state.last_save_time = datetime.now()
                st.rerun()

        if st.session_state.get('avaliacao_iniciada', False):
            if 'last_save_time' not in st.session_state:
                st.session_state.last_save_time = datetime.now()
            if datetime.now() - st.session_state.last_save_time > timedelta(minutes=10):
                try:
                    with open(st.session_state.caminho_arquivo, 'w', encoding='utf-8') as f: 
                        json.dump(st.session_state.respostas, f, ensure_ascii=False, indent=4)
                    st.session_state.last_save_time = datetime.now()
                    st.toast(f"Progresso salvo automaticamente às {datetime.now().strftime('%H:%M:%S')}")
                except Exception as e: 
                    st.toast(f"Erro no salvamento automático: {e}")
            
            st.header(f"Avaliação: {st.session_state.municipio} - {st.session_state.segmento}")
            
            if st.session_state.segmento not in matriz_completa:
                st.error("Dados de perguntas não encontrados para o segmento selecionado.")
                st.stop()

            matriz_perguntas_segmento = matriz_completa[st.session_state.segmento]
            
            current_results = calcular_indice_e_selo(st.session_state.respostas, matriz_perguntas_segmento)
            st.info(f"**Índice Geral de Transparência:** {current_results['indice']:.2f}% | **Selo Atricon:** {current_results['selo']}")

            for secao, perguntas in matriz_perguntas_segmento.items():
                if secao == "Municipios_MA": continue
                with st.expander(f"**{secao}**", expanded=False):
                    for item in perguntas:
                        st.markdown(f"#### {item['topico']} - {item['criterio']}"); st.markdown("---")
                        
                        col_link_ui, _ = st.columns([1, 1])
                        with col_link_ui:
                            st.subheader("Links de Evidência")
                            chave_links = f"{secao}_{item['criterio']}_links"
                            if chave_links not in st.session_state.respostas:
                                st.session_state.respostas[chave_links] = []
                            
                            for i, link in enumerate(st.session_state.respostas[chave_links]):
                                link_cols = st.columns([10, 1])
                                link_cols[0].info(link)
                                if link_cols[1].button("X", key=f"rem_{chave_links}_{i}"): 
                                    st.session_state.respostas[chave_links].pop(i)
                                    st.rerun()
                            
                            link_cols = st.columns([10, 1])
                            novo_link_key = f"add_{chave_links}"
                            novo_link = link_cols[0].text_input("Adicionar novo link", value="", key=novo_link_key, label_visibility="collapsed")
                            if link_cols[1].button("➕", key=f"btn_{chave_links}"):
                                if novo_link and novo_link not in st.session_state.respostas[chave_links]:
                                    st.session_state.respostas[chave_links].append(novo_link)
                                    st.rerun()


                        st.markdown("---"); st.subheader("Critérios de Avaliação")
                        subcriterios = item["subcriterios"]
                        
                        disponibilidade_falhou_na_sessao = False
                        if "Disponibilidade" in subcriterios:
                            chave_disponibilidade_resposta = f"{secao}_{item['criterio']}_Disponibilidade"
                            if st.session_state.respostas.get(chave_disponibilidade_resposta) == "Não Atende":
                                disponibilidade_falhou_na_sessao = True 

                            cols = st.columns([1, 2])
                            with cols[0]:
                                resposta_atual_disp = st.session_state.respostas.get(chave_disponibilidade_resposta, "Atende")
                                st.radio("Disponibilidade", ("Atende", "Não Atende"), 
                                         index=1 if resposta_atual_disp == "Não Atende" else 0, 
                                         key=chave_disponibilidade_resposta, 
                                         horizontal=True, 
                                         on_change=on_disponibilidade_change, 
                                         kwargs=dict(secao=secao, criterio=item['criterio'], subcriterios=subcriterios))
                            
                            if st.session_state.respostas.get(chave_disponibilidade_resposta) == "Não Atende":
                                with cols[1]:
                                    chave_obs_disp = f"{chave_disponibilidade_resposta}_obs"
                                    obs_disp = st.text_area("Observação:", value=st.session_state.respostas.get(chave_obs_disp, ""), key=chave_obs_disp)
                                    st.session_state.respostas[chave_obs_disp] = obs_disp
                                    
                        for subcriterio in subcriterios:
                            if subcriterio == "Disponibilidade":
                                continue
                            
                            cols = st.columns([1, 2])
                            chave_resposta_sub = f"{secao}_{item['criterio']}_{subcriterio}"
                            
                            with cols[0]:
                                resposta_atual_sub = st.session_state.respostas.get(chave_resposta_sub, "Atende")
                                
                                disabled = disponibilidade_falhou_na_sessao
                                display_index = 1 if resposta_atual_sub == "Não Atende" else 0

                                st.radio(subcriterio, ("Atende", "Não Atende"), 
                                         index=display_index, 
                                         key=chave_resposta_sub, 
                                         horizontal=True, 
                                         disabled=disabled)
                                st.session_state.respostas[chave_resposta_sub] = ["Atende", "Não Atende"][display_index]
                                
                            if st.session_state.respostas.get(chave_resposta_sub) == "Não Atende":
                                with cols[1]:
                                    chave_obs_sub = f"{chave_resposta_sub}_obs"
                                    obs_sub = st.text_area("Observação:", value=st.session_state.respostas.get(chave_obs_sub, ""), key=chave_obs_sub, disabled=disabled)
                                    st.session_state.respostas[chave_obs_sub] = obs_sub
                                    
                        st.markdown("---")
            
            st.sidebar.header("Ações")
            if st.sidebar.button("💾 Salvar Progresso"):
                try:
                    with open(st.session_state.caminho_arquivo, 'w', encoding='utf-8') as f: 
                        json.dump(st.session_state.respostas, f, ensure_ascii=False, indent=4)
                    st.session_state.last_save_time = datetime.now()
                    st.sidebar.success("Progresso salvo!")
                except Exception as e:
                    st.sidebar.error(f"Erro ao salvar progresso: {e}")

            st.sidebar.markdown("##### Tipo de Relatório")
            tipo_relatorio = st.sidebar.radio("Escolha o tipo:", ("Apenas Não Conformidades", "Relatório Completo"), label_visibility="collapsed")
            
            if st.sidebar.button("📊 Gerar Relatório PDF"):
                with st.spinner("Gerando relatório PDF..."):
                    try:
                        with open(st.session_state.caminho_arquivo, 'w', encoding='utf-8') as f: 
                            json.dump(st.session_state.respostas, f, ensure_ascii=False, indent=4)
                        st.session_state.last_save_time = datetime.now()
                    except Exception as e:
                        st.error(f"Erro ao salvar antes de gerar relatório: {e}")
                        st.stop() 

                    st.session_state.path_pdf = None
                    st.session_state.fallback_docx_path = None
                    
                    docx_output_path, pdf_output_path = gerar_relatorio_novo_modelo(
                        st.session_state.respostas, 
                        st.session_state.municipio, 
                        st.session_state.segmento, 
                        matriz_completa[st.session_state.segmento], 
                        tipo_relatorio, 
                        st.session_state["name"], 
                        config['credentials']['usernames'][st.session_state['username']]
                    )
                    
                    st.session_state.path_pdf = pdf_output_path
                    st.session_state.fallback_docx_path = docx_output_path 
                    
                if st.session_state.path_pdf:
                    st.sidebar.success("Relatório PDF pronto para download!")
                elif st.session_state.fallback_docx_path:
                    st.sidebar.warning("Falha ao gerar PDF, mas o arquivo DOCX está pronto para download.")
                
                st.rerun()

            if st.session_state.get('path_pdf'):
                with open(st.session_state.path_pdf, "rb") as pdf_file:
                    st.sidebar.download_button(
                        label="⬇️ Baixar Relatório (.pdf)", 
                        data=pdf_file, 
                        file_name=os.path.basename(st.session_state.path_pdf), 
                        mime="application/pdf", 
                        key="download_pdf"
                    )
            
            if st.session_state.get('fallback_docx_path'):
                with open(st.session_state.fallback_docx_path, "rb") as docx_file:
                    st.sidebar.download_button(
                        label="⬇️ Baixar Arquivo Word (.docx)", 
                        data=docx_file, 
                        file_name=os.path.basename(st.session_state.fallback_docx_path), 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                        key="download_docx_fallback"
                    )

    elif st.session_state["authentication_status"] is False:
        st.error('Usuário ou senha incorretos.')
    elif st.session_state["authentication_status"] is None:
        st.warning('Por favor, insira seu usuário e senha para acessar o sistema.')
else:
    st.warning("Aguardando o carregamento do arquivo 'criterios_por_topico.json'... Verifique se ele existe e está formatado corretamente.")